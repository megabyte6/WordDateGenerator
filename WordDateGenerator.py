#!/usr/bin/env python

import os.path
import webbrowser
from datetime import date, datetime, timedelta
from typing import Callable

import customtkinter as ctk
from docx import Document
from tkcalendar import DateEntry


class DateRangeElement(ctk.CTkFrame):
    def __init__(
        self,
        master: ctk.CTkFrame,
        start_date: datetime = date.today(),
        end_date: datetime = date.today(),
        on_remove: Callable[[ctk.CTkFrame], None] = None,
        *args,
        **kwargs,
    ):
        super().__init__(master, *args, **kwargs)

        self._on_remove = on_remove

        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        self._ui_start_date_entry = DateEntry(self)
        self._ui_start_date_entry.set_date(start_date)
        self._ui_start_date_entry.grid(row=0, column=0, padx=(0, 5))

        self._ui_end_date_entry = DateEntry(self)
        self._ui_end_date_entry.set_date(end_date)
        self._ui_end_date_entry.grid(row=0, column=1, padx=5)

        self._ui_remove_button = ctk.CTkButton(self, width=30, height=30, text="x", command=self._remove_from_list)
        self._ui_remove_button.grid(row=0, column=2, padx=(5, 0))

    def _remove_from_list(self):
        if self._on_remove:
            self._on_remove(self)

        self.destroy()

    def enabled(self, enabled: bool) -> None:
        if enabled:
            self._ui_start_date_entry.configure(state="normal")
            self._ui_end_date_entry.configure(state="normal")
            self._ui_remove_button.configure(state="normal")
        else:
            self._ui_start_date_entry.configure(state="disabled")
            self._ui_end_date_entry.configure(state="disabled")
            self._ui_remove_button.configure(state="disabled")

    def get_start_date(self) -> datetime:
        return self._ui_start_date_entry.get_date()

    def get_end_date(self) -> datetime:
        return self._ui_end_date_entry.get_date()


class WordDateGenerator:
    def __init__(
        self,
        path: str,
        start_date: datetime = None,
        end_date: datetime = None,
        date_format: str = "%a. %b. %d",
        excluded_days: set[str] = None,
        excluded_date_ranges: list[DateRangeElement] = None,
        date_column: int = 0,
    ) -> None:
        if not os.path.exists(path) or not os.path.isfile(path):
            raise FileNotFoundError(f"File {path} does not exist.")
        self.path = path

        if start_date is None:
            start_date = date.today()
        self.start_date = start_date

        if end_date is None:
            end_date = date.today() + timedelta(days=7)
        self.end_date = end_date

        self.date_format = date_format

        if excluded_days is None:
            excluded_days = {"saturday", "sunday"}
        self.excluded_days = excluded_days

        if excluded_date_ranges is None:
            excluded_date_ranges = []
        self.excluded_date_ranges = excluded_date_ranges

        self.date_column = date_column

        self.doc = Document(path)

        self.selected_table = self.doc.tables[0]

    def _generate_dates(self, start_date: datetime, end_date: datetime, excluded_days: list[str]) -> list[datetime]:
        """
        Generates a list of dates between start_date and end_date excluding weekends.

        Args:
            start_date (datetime): The start date. Defaults to today.
            end_date (datetime): The end date. Defaults to one week from today.
            excluded_days (list[str]): A list of days to exclude. Defaults to ["saturday", "sunday"].

        Returns:
            list[datetime]: A list of dates between start_date and end_date excluding weekends.
        """

        days_as_numbers = {
            "monday": 0,
            "tuesday": 1,
            "wednesday": 2,
            "thursday": 3,
            "friday": 4,
            "saturday": 5,
            "sunday": 6,
        }
        excluded_numbers = [days_as_numbers[day.lower()] for day in excluded_days]
        excluded_dates = {
            current_date
            for date_range_element in self.excluded_date_ranges
            for current_date in (
                date_range_element.get_start_date() + timedelta(days=i)
                for i in range((date_range_element.get_end_date() - date_range_element.get_start_date()).days + 1)
            )
        }

        return [
            start_date + timedelta(days=i)
            for i in range((end_date - start_date).days + 1)
            if (start_date + timedelta(days=i)).weekday() not in excluded_numbers
            and (start_date + timedelta(days=i)) not in excluded_dates
        ]

    def add_dates_to_table(self) -> None:
        """
        Adds dates to the table by iterating through each row and assigning a date value to the first cell.
        """

        dates = self._generate_dates(self.start_date, self.end_date, self.excluded_days)
        for row, date in zip(self.selected_table.rows, dates):
            row.cells[self.date_column].text = date.strftime(self.date_format)

    def save(self, path: str = None) -> None:
        """
        Save the document.

        Args:
            path (str, optional): The path to save the document to. Will overwrite the existing file by default.
        """

        if path is None:
            path = self.path

        self.doc.save(path)


class App(ctk.CTk):
    def __init__(self) -> None:
        super().__init__()

        self.title("Word Date Generator")
        self.geometry("900x550")

        path_frame = ctk.CTkFrame(self, fg_color="transparent")
        path_frame.pack(fill=ctk.X, padx=20, pady=20)
        path_frame.grid_columnconfigure(0, weight=1)

        self._ui_path_entry = ctk.CTkEntry(path_frame, height=50, placeholder_text="Enter the path to the document")
        self._ui_path_entry.grid(row=0, column=0, sticky="EW")
        self._ui_path_entry.bind("<KeyRelease>", self._path_entry_handler)

        self._ui_path_picker = ctk.CTkButton(
            path_frame, height=50, width=75, text="Choose", command=self._open_file_picker
        )
        self._ui_path_picker.grid(row=0, column=1)

        options_frame = ctk.CTkFrame(self, fg_color="transparent")
        options_frame.pack()

        start_date_label = ctk.CTkLabel(options_frame, text="Start Date:")
        start_date_label.grid(row=0, column=0, padx=5, pady=10, sticky="E")

        self._ui_start_date_picker = DateEntry(options_frame)
        self._ui_start_date_picker.grid(row=0, column=1, padx=5, pady=10)
        self._ui_start_date_picker.bind("<<DateEntrySelected>>", self._start_date_picker_handler)
        self._ui_start_date_picker.bind("<Return>", self._start_date_picker_handler)
        self._ui_start_date_picker.bind("<FocusOut>", self._start_date_picker_handler)
        self._ui_start_date_picker.bind("<Leave>", self._start_date_picker_handler)

        end_date_label = ctk.CTkLabel(options_frame, text="End Date:")
        end_date_label.grid(row=1, column=0, padx=5, pady=10, sticky="E")

        end_date = date.today() + timedelta(days=7)
        self._ui_end_date_picker = DateEntry(options_frame, year=end_date.year, month=end_date.month, day=end_date.day)
        self._ui_end_date_picker.grid(row=1, column=1, padx=5, pady=10)
        self._ui_end_date_picker.bind("<<DateEntrySelected>>", self._end_date_picker_handler)
        self._ui_end_date_picker.bind("<Return>", self._end_date_picker_handler)
        self._ui_end_date_picker.bind("<FocusOut>", self._end_date_picker_handler)
        self._ui_end_date_picker.bind("<Leave>", self._end_date_picker_handler)

        table_index_label = ctk.CTkLabel(options_frame, text="Table:")
        table_index_label.grid(row=2, column=0, padx=5, pady=10, sticky="E")

        self._ui_table_index = ctk.CTkComboBox(options_frame, values=["1"], command=self._table_index_handler)
        self._ui_table_index.grid(row=2, column=1, padx=5, pady=10)

        table_column_label = ctk.CTkLabel(options_frame, text="Date Column:")
        table_column_label.grid(row=3, column=0, padx=5, pady=10, sticky="E")

        self._ui_table_column = ctk.CTkComboBox(options_frame, values=["1"], command=self._table_column_handler)
        self._ui_table_column.grid(row=3, column=1, padx=5, pady=10)

        date_format_label = ctk.CTkLabel(options_frame, text="Date Format:")
        date_format_label.grid(row=4, column=0, padx=5, pady=10, sticky="E")

        date_format_frame = ctk.CTkFrame(options_frame, fg_color="transparent")
        date_format_frame.grid(row=4, column=1, padx=5, pady=10)

        self._ui_date_format = ctk.CTkEntry(date_format_frame)
        self._ui_date_format.insert(0, "%a. %b. %d")
        self._ui_date_format.grid(row=0, column=0, padx=5)
        self._ui_date_format.bind("<KeyRelease>", self._date_format_handler)

        date_format_info_label = ctk.CTkLabel(date_format_frame, text="?")
        date_format_info_label.grid(row=0, column=1)
        date_format_info_label.bind(
            "<Button>",
            lambda e: webbrowser.open(
                "https://docs.python.org/3/library/datetime.html#strftime-and-strptime-format-codes"
            ),
        )

        self._ui_date_format_preview = ctk.CTkLabel(
            date_format_frame, text=date.today().strftime(self._ui_date_format.get())
        )
        self._ui_date_format_preview.grid(row=1, column=0, columnspan=2, padx=5)

        exclude_days_frame = ctk.CTkFrame(options_frame)
        exclude_days_frame.grid(row=0, column=2, rowspan=options_frame.grid_size()[1], padx=30)

        exclude_days_label = ctk.CTkLabel(exclude_days_frame, text="Days to exclude:")
        exclude_days_label.grid(row=0, padx=15, pady=(10, 5))

        self.exclude_day_checkboxes: list[ctk.CTkCheckBox] = []

        exclude_monday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame, text="Monday", command=lambda: self._weekday_checkbox_handler(exclude_monday_checkbox)
        )
        exclude_monday_checkbox.grid(row=1, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_monday_checkbox)

        exclude_tuesday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame, text="Tuesday", command=lambda: self._weekday_checkbox_handler(exclude_tuesday_checkbox)
        )
        exclude_tuesday_checkbox.grid(row=2, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_tuesday_checkbox)

        exclude_wednesday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame,
            text="Wednesday",
            command=lambda: self._weekday_checkbox_handler(exclude_wednesday_checkbox),
        )
        exclude_wednesday_checkbox.grid(row=3, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_wednesday_checkbox)

        exclude_thursday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame,
            text="Thursday",
            command=lambda: self._weekday_checkbox_handler(exclude_thursday_checkbox),
        )
        exclude_thursday_checkbox.grid(row=4, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_thursday_checkbox)

        exclude_friday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame, text="Friday", command=lambda: self._weekday_checkbox_handler(exclude_friday_checkbox)
        )
        exclude_friday_checkbox.grid(row=5, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_friday_checkbox)

        exclude_saturday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame,
            text="Saturday",
            command=lambda: self._weekday_checkbox_handler(exclude_saturday_checkbox),
            variable=ctk.BooleanVar(self, value=True),
        )
        exclude_saturday_checkbox.grid(row=6, padx=15, pady=5, sticky="W")
        self.exclude_day_checkboxes.append(exclude_saturday_checkbox)

        exclude_sunday_checkbox = ctk.CTkCheckBox(
            exclude_days_frame,
            text="Sunday",
            command=lambda: self._weekday_checkbox_handler(exclude_sunday_checkbox),
            variable=ctk.BooleanVar(self, value=True),
        )
        exclude_sunday_checkbox.grid(row=7, padx=15, pady=(5, 15), sticky="W")
        self.exclude_day_checkboxes.append(exclude_sunday_checkbox)

        self._ui_exclude_range_frame = ctk.CTkFrame(options_frame)
        self._ui_exclude_range_frame.grid(row=0, column=4, rowspan=options_frame.grid_size()[1], padx=10, sticky="N")

        exclude_range_label = ctk.CTkLabel(self._ui_exclude_range_frame, text="Exclude Range:")
        exclude_range_label.grid(row=0, padx=15, pady=(10, 5))

        self.exclude_ranges: list[DateRangeElement] = []

        self._ui_exclude_new_range = ctk.CTkButton(
            self._ui_exclude_range_frame, text="+", command=self._exclude_new_range_handler
        )
        self._ui_exclude_new_range.grid(row=1, padx=15, pady=(10, 15), sticky="EW")

        self._ui_save_as_new_file = ctk.CTkCheckBox(
            self, text="Save as new file", variable=ctk.BooleanVar(self, value=True)
        )
        self._ui_save_as_new_file.pack(pady=(30, 20))

        self._ui_generate_button = ctk.CTkButton(self, height=50, width=150, text="Generate", command=self._generate)
        self._ui_generate_button.pack()

        self._set_ui_state(enabled=False)

    def _set_ui_state(self, enabled: bool) -> None:
        """
        Sets the UI state based on the value of `enabled`.

        Args:
            enabled (bool): A boolean value indicating whether the UI should be enabled or disabled.

        Returns:
            None: This function does not return anything.
        """

        if enabled:
            self._ui_start_date_picker.configure(state="normal")
            self._ui_end_date_picker.configure(state="normal")
            self._ui_table_index.configure(state="readonly")
            self._ui_table_column.configure(state="readonly")
            self._ui_date_format.configure(state="normal")
            for checkbox in self.exclude_day_checkboxes:
                checkbox.configure(state="normal")
            self._ui_exclude_new_range.configure(state="normal")
            for range_element in self.exclude_ranges:
                range_element.enabled(True)
            self._ui_save_as_new_file.configure(state="normal")
            self._ui_generate_button.configure(state="normal")
        else:
            self._ui_start_date_picker.configure(state="disabled")
            self._ui_end_date_picker.configure(state="disabled")
            self._ui_table_index.configure(state="disabled")
            self._ui_table_column.configure(state="disabled")
            self._ui_date_format.configure(state="disabled")
            for checkbox in self.exclude_day_checkboxes:
                checkbox.configure(state="disabled")
            self._ui_exclude_new_range.configure(state="disabled")
            for range_element in self.exclude_ranges:
                range_element.enabled(False)
            self._ui_save_as_new_file.configure(state="disabled")
            self._ui_generate_button.configure(state="disabled")

    def _open_file_picker(self) -> None:
        """
        Opens a file picker and creates the document object with the chosen file.
        """

        path = ctk.filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            self._ui_path_entry.delete(0, ctk.END)
            self._ui_path_entry.insert(0, path)
            self._path_entry_handler()

    def _path_entry_handler(self, *args) -> None:
        try:
            self.document = WordDateGenerator(self._ui_path_entry.get())
            self._set_ui_state(enabled=True)
        except FileNotFoundError:
            self._set_ui_state(enabled=False)
            self._ui_generate_button.configure(text="Generate")
            return

        table_options = [str(i + 1) for i in range(len(self.document.doc.tables))]
        self._ui_table_index.configure(values=table_options)
        self._ui_table_index.set(table_options[0])

        # Write options showing on-screen to the WordDateGenerator object
        self.document.start_date = self._ui_start_date_picker.get_date()
        self.document.end_date = self._ui_end_date_picker.get_date()
        self.document.selected_table = self.document.doc.tables[int(self._ui_table_index.get()) - 1]
        self.document.date_column = int(self._ui_table_column.get()) - 1
        self.document.date_format = self._ui_date_format.get()
        self.document.excluded_days = {
            checkbox.cget("text").lower() for checkbox in self.exclude_day_checkboxes if checkbox.get()
        }
        self.document.excluded_date_ranges = self.exclude_ranges

    def _start_date_picker_handler(self, *args) -> None:
        if hasattr(self, "document"):
            self.document.start_date = self._ui_start_date_picker.get_date()

    def _end_date_picker_handler(self, *args) -> None:
        if hasattr(self, "document"):
            self.document.end_date = self._ui_end_date_picker.get_date()

    def _table_index_handler(self, *args) -> None:
        self.document.selected_table = self.document.doc.tables[int(self._ui_table_index.get()) - 1]

        column_options = [str(i + 1) for i in range(len(self.document.selected_table.columns))]
        self._ui_table_column.configure(values=column_options)
        self._ui_table_column.set(column_options[0])

    def _table_column_handler(self, *args) -> None:
        self.document.date_column = int(self._ui_table_column.get()) - 1

    def _weekday_checkbox_handler(self, weekday_checkbox: ctk.CTkCheckBox) -> None:
        checkbox_name = weekday_checkbox.cget("text").lower()
        if weekday_checkbox.get():
            self.document.excluded_days.add(checkbox_name)
        else:
            self.document.excluded_days.remove(checkbox_name)

    def _date_format_handler(self, *args) -> None:
        date_format = self._ui_date_format.get()
        self.document.date_format = date_format
        self._ui_date_format_preview.configure(text=date.today().strftime(date_format))

    def _exclude_new_range_handler(self) -> None:
        date_range = DateRangeElement(self._ui_exclude_range_frame, on_remove=self._remove_date_range_element)
        self.exclude_ranges.append(date_range)
        self.document.excluded_date_ranges.append(date_range)

        current_row = self._ui_exclude_new_range.grid_info()["row"]
        self._ui_exclude_new_range.grid(row=current_row + 1)
        date_range.grid(row=current_row, padx=15, pady=10)

    def _remove_date_range_element(self, date_range_element: DateRangeElement) -> None:
        self.exclude_ranges.remove(date_range_element)
        self.document.excluded_date_ranges.remove(date_range_element)

    def _generate(self) -> None:
        """
        Fills the table with dates and saves the file.
        """

        self.document.add_dates_to_table()

        if self._ui_save_as_new_file.get():
            path = ctk.filedialog.asksaveasfilename(filetypes=[("Word Documents", "*.docx")])
            if path:
                if not path.endswith(".docx"):
                    path += ".docx"
                self.document.save(path)
            else:
                return
        else:
            self.document.save()

        self._ui_generate_button.configure(text="Done!")


if __name__ == "__main__":
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")

    app = App()
    app.mainloop()
